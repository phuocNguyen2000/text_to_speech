import time
from gtts import gTTS
import os
from docx import Document

# Input text you want to convert to Vietnamese speech
input_docx = "input/comatulau.docx"
output_directory = "out/"

# Create the output directory if it doesn't exist
import os
# if not os.path.exists(output_directory):
#     os.makedirs(output_directory)

# # Load the input DOCX file
# file_ref = open("comatulau.docx")
# doc = Document(docx=input_docx)

# # Initialize variables for page tracking
# current_page = 0
# file_index = 1
# current_doc = Document()
# pages_per_subdocument = 10

# # Initialize variables
# current_page = 0
# subdoc_number = 1

# # Iterate through the paragraphs to split the document
# for paragraph in doc.paragraphs:
#     current_page += 1
#     current_doc.add_paragraph(paragraph.text)

#     # Check if we've reached 10 pages
#     if current_page % 10 == 0:
#         # Save the current document to a new output file
#         output_filename = f"{output_directory}output_file_{file_index}.docx"
#         current_doc.save(output_filename)
#         print(f"Saved {output_filename}")

#         # Create a new document for the next batch of pages
#         current_doc = Document()
#         file_index += 1

# # If there are remaining pages, save them to the last file
# if len(current_doc.paragraphs) > 0:
#     output_filename = f"{output_directory}output_file_{file_index}.docx"
#     current_doc.save(output_filename)
#     print(f"Saved {output_filename}")
    
for filename in os.listdir('out/1-10/'):
    f = os.path.join('out/1-10/', filename)
    # checking if it is a file
    if os.path.isfile(f):
        print(filename)
    s=str(f)
    doc = Document(docx=s)
    text_to_convert = ""
    for paragraph in doc.paragraphs:
        text_to_convert=text_to_convert+ " " +paragraph.text
    
    

    # Create a gTTS object with the Vietnamese language code "vi"
    tts = gTTS(text=text_to_convert, lang="vi")
    s=s.replace('\\', '')
    s=s.replace('out/1-10/', '')
    s=s.replace('.docx', '')
    # Save the speech as an audio file (e.g., output.mp3)
    tts.save(f"out_mp3/{s}.mp3")
    time.sleep(10)

    # Play the audio file (on systems with a default audio player)
    # os.system("start output.mp3")  # On Windows
    # On Linux/Mac, you can use a player like VLC: os.system("vlc output.mp3")